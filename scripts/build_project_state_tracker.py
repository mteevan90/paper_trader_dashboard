"""build_project_state_tracker.py — generate docs/Project_State_Tracker.docx.

What it is: serializes the canonical project-state tracker content (audit +
strategic context + cleanup summary) into a Word document for the repo. The
source text lives inside this file — same pattern as build_project_state_audit.py
— so the script is self-contained and reproducible without external state.

When to re-run: any time the project state shifts enough that the existing
tracker is stale. Update the literal strings in this file, then re-run.

How to invoke (PowerShell):
    venv\\Scripts\\python.exe scripts\\build_project_state_tracker.py

Output: docs/Project_State_Tracker.docx (committed to repo).

Formatting conventions:
- Arial body 11pt; Arial Bold for headings; blue (#1F4E79) heading color.
- US Letter page size.
- Status table rows color-coded per [RED]/[YELLOW]/[GREEN] tags in the source.
- Code blocks in Consolas 9pt with light-gray cell shading.

Requirements: python-docx (`pip install python-docx`). Not in
requirements.txt — install ad-hoc if needed for a tracker refresh.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = REPO_ROOT / "docs" / "Project_State_Tracker.docx"

BODY_FONT = "Arial"
CODE_FONT = "Consolas"
HEADING_BLUE = RGBColor(0x1F, 0x4E, 0x79)
TEXT_BLACK = RGBColor(0x00, 0x00, 0x00)
TEXT_GRAY = RGBColor(0x50, 0x50, 0x50)
CODE_BG = "F2F2F2"     # very light gray
ONBOARDING_BG = "F5F5F5"  # very light gray (distinct from code blocks)
STATUS_RED = "F4CCCC"  # light red
STATUS_YELLOW = "FFF2CC"  # light yellow
STATUS_GREEN = "D5E8D4"  # light green
HEADER_BG = "1F4E79"   # dark blue (status table header)
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Low-level XML helpers (python-docx doesn't expose cell shading directly)
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


# ---------------------------------------------------------------------------
# Style + paragraph helpers
# ---------------------------------------------------------------------------

def configure_document(doc: Document) -> None:
    # US Letter, portrait, 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_BLACK

    for level in (1, 2, 3):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = BODY_FONT
            s.font.bold = True
            s.font.color.rgb = HEADING_BLUE
            s.font.size = Pt(20 - 3 * (level - 1))   # 20, 17, 14


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def para(doc, text, *, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    return p


def para_runs(doc, parts):
    """Add a paragraph from a list of (text, bold, italic) tuples for
    inline formatting."""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.italic = italic
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    return p


def bullet_runs(doc, parts):
    """Bulleted paragraph with mixed-formatting runs."""
    p = doc.add_paragraph(style="List Bullet")
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.italic = italic
    return p


def code_block(doc, code_text: str, *, bg_color: str = CODE_BG) -> None:
    """Render a code block as a single-cell table with gray fill + Consolas."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, bg_color)
    cell.text = ""    # clear default
    for line in code_text.split("\n"):
        p = cell.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = CODE_FONT
        r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    # Remove the auto-added empty first paragraph (cell.text="" leaves one
    # blank paragraph; we appended N more, so total is N+1).
    first_p = cell.paragraphs[0]
    if not first_p.text:
        p_elem = first_p._element
        p_elem.getparent().remove(p_elem)


def divider(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run("⸻" * 1)
    r.font.color.rgb = TEXT_GRAY


# ---------------------------------------------------------------------------
# Embedded onboarding prompt
# ---------------------------------------------------------------------------

ONBOARDING_PROMPT = (
    "I'm working on a quant equity / options paper trading project at\n"
    "github.com/mteevan90/paper_trader_dashboard. Two contributors: Mike\n"
    "(equity research; account @mteevan90) and Chris (options + paused\n"
    "crypto; @cmjteevan).\n"
    "\n"
    "Before answering my first question, please read these in order:\n"
    "\n"
    "1. docs/Project_State_Tracker.docx — canonical current state, action\n"
    "   items, strategic context. THIS is your primary orientation doc.\n"
    "\n"
    "2. docs/Comprehensive_User_Guide.docx — full technical reference\n"
    "   (architecture, data layer, dashboard patterns, gotchas). Reference\n"
    "   for depth.\n"
    "\n"
    "3. docs/Project_State_Audit.docx — most recent audit from Claude Code\n"
    "   (2026-05-11), with file-level details on what's broken / stale /\n"
    "   clean.\n"
    "\n"
    "4. docs/Options_Extension_Decisions.md — Chris's options architecture\n"
    "   decisions.\n"
    "\n"
    "5. docs/Crypto_Extension_Decisions.md — Crypto extension architectural\n"
    "   decisions (extension is paused but doc is authoritative if work\n"
    "   resumes).\n"
    "\n"
    "6. docs/future_work.md — deferred items, including hardware-aware work\n"
    "   allocation and distributed Optuna trigger.\n"
    "\n"
    "Once you've read these, give me a 3-sentence summary of where the\n"
    "project stands and what URGENT items remain, then ask what I want to\n"
    "work on. Don't write any code or run any tools until I respond. Don't\n"
    "make assumptions about my role (Mike or Chris) — ask if it's not\n"
    "clear from context.\n"
    "\n"
    "Key project conventions to absorb:\n"
    "- Mike owns equity work + shared infrastructure\n"
    "- Chris owns options + crypto code\n"
    "- CODEOWNERS in .github/ enforces cross-namespace boundaries\n"
    "- Cloud dashboard at paper-trader-mteev.streamlit.app (read-only,\n"
    "  R2-backed)\n"
    "- Asset-class-namespaced data layer: equities/, crypto/, options/\n"
    "  subdirs under models/cache/ and models/snapshots/\n"
    "- The tracker doc supersedes prior versions of any other handoff doc\n"
    "\n"
    "Three things you should specifically NOT do without explicit\n"
    "confirmation:\n"
    "- Don't push to main (use feature branches; let me merge via PR)\n"
    "- Don't run snapshot_for_cloud.py without --dry-run first\n"
    "- Don't pop or modify stash@{0} (deferred SP1500 work)\n"
    "\n"
    "If you find yourself working without having read the tracker doc\n"
    "first, stop and read it. The tracker is updated when project state\n"
    "changes meaningfully — its commit date in git log tells you how\n"
    "recently."
)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

# (workstream, status, owner, blocker_or_next, color_tag)
STATUS_ROWS = [
    ("Equity research framework",
     "WORKING",
     "Mike",
     "None; dashboard updates needed for NVDA/META finding",
     "GREEN"),
    ("Cloud dashboard (equities)",
     "WORKING but R2 stale",
     "Mike",
     "Sync R2 after PR #15 merges",
     "YELLOW"),
    ("Options Sections 1-8 + 2.5",
     "MERGED",
     "Chris",
     "Done; foundation for v1 study",
     "GREEN"),
    ("Options PR #15 (hang fix)",
     "OPEN — URGENT",
     "Chris",
     "Mike review + merge",
     "RED"),
    ("Options v1 study",
     "BLOCKED on PR #15",
     "Chris",
     "Cannot run until merge",
     "RED"),
    ("Crypto Phase 2 strategy code",
     "PAUSED — STUB ONLY",
     "Chris",
     "Indefinitely; Chris on options",
     "YELLOW"),
    ("SP1500 expansion",
     "PAUSED — STASHED",
     "Mike",
     "Finnhub TOS OR pivot to Polygon (see §3)",
     "YELLOW"),
    ("NVDA/META dashboard surfacing",
     "QUEUED",
     "Mike",
     "Small Claude Code session; commits ready (f46bd05)",
     None),
    ("Test suite",
     "374 pass / 2 venv-fixable",
     "Shared",
     "Run pip install -r requirements.txt",
     "YELLOW"),
    ("Repo hygiene",
     "CLEANED 2026-05-11",
     "Mike",
     "Commit 5c39cc5",
     "GREEN"),
]


def render_status_table(doc):
    table = doc.add_table(rows=1 + len(STATUS_ROWS), cols=4)
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(1.7)
    table.columns[2].width = Inches(0.7)
    table.columns[3].width = Inches(2.2)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(("Workstream", "Status", "Owner", "Blocker / Next")):
        _set_cell_shading(hdr_cells[i], HEADER_BG)
        _set_cell_borders(hdr_cells[i])
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = HEADER_FG
        r.font.name = BODY_FONT
        r.font.size = Pt(10)

    color_map = {"RED": STATUS_RED, "YELLOW": STATUS_YELLOW,
                 "GREEN": STATUS_GREEN}
    for row_idx, (ws, st, ow, bn, color) in enumerate(STATUS_ROWS, 1):
        cells = table.rows[row_idx].cells
        if color in color_map:
            for c in cells:
                _set_cell_shading(c, color_map[color])
        for c, v in zip(cells, (ws, st, ow, bn)):
            _set_cell_borders(c)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(v)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)


POLYGON_TIERS = [
    ("Free (Basic)", "$0", "2 years EOD", "5 calls/min", "No (15-min delayed)"),
    ("Starter", "~$29", "5 years EOD", "Unlimited", "No (15-min delayed)"),
    ("Developer", "~$79", "10+ years EOD + intraday", "Unlimited",
     "Real-time included"),
    ("Advanced", "~$199", "20+ years + tick data", "Unlimited",
     "Real-time + WebSocket"),
]


def render_polygon_table(doc):
    table = doc.add_table(rows=1 + len(POLYGON_TIERS), cols=5)
    table.autofit = False
    table.columns[0].width = Inches(1.1)
    table.columns[1].width = Inches(0.8)
    table.columns[2].width = Inches(1.8)
    table.columns[3].width = Inches(1.1)
    table.columns[4].width = Inches(1.7)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(
            ("Tier", "Price/mo", "Historical depth",
             "Rate limit", "Real-time?")):
        _set_cell_shading(hdr_cells[i], HEADER_BG)
        _set_cell_borders(hdr_cells[i])
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = HEADER_FG
        r.font.name = BODY_FONT
        r.font.size = Pt(10)
    for row_idx, row in enumerate(POLYGON_TIERS, 1):
        cells = table.rows[row_idx].cells
        for c, v in zip(cells, row):
            _set_cell_borders(c)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(v)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)


LINEAGE_ROWS = [
    ("Paper_Trader_Handoff.docx (v1.0)", "2026-05-08",
     "Initial handoff to Chris before crypto extension"),
    ("Crypto_Extension_Decisions.md", "2026-05-09",
     "Architectural decisions for multi-asset support"),
    ("Comprehensive_User_Guide.docx (v2.0)", "2026-05-09",
     "Full user guide post-Phase-1; NVDA/META findings"),
    ("Options_Extension_Decisions.md", "2026-05-10",
     "Architectural decisions for options extension"),
    ("Project_State_Audit.docx", "2026-05-11",
     "Claude Code audit findings (raw audit)"),
    ("Project_State_Tracker.docx (this doc)", "2026-05-11",
     "Audit + strategic context + cleanup summary"),
]


def render_lineage_table(doc):
    table = doc.add_table(rows=1 + len(LINEAGE_ROWS), cols=3)
    table.autofit = False
    table.columns[0].width = Inches(2.6)
    table.columns[1].width = Inches(1.0)
    table.columns[2].width = Inches(2.9)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(("Document", "Date", "Purpose")):
        _set_cell_shading(hdr_cells[i], HEADER_BG)
        _set_cell_borders(hdr_cells[i])
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = HEADER_FG
        r.font.name = BODY_FONT
        r.font.size = Pt(10)
    for row_idx, row in enumerate(LINEAGE_ROWS, 1):
        cells = table.rows[row_idx].cells
        for c, v in zip(cells, row):
            _set_cell_borders(c)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(v)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Main document content
# ---------------------------------------------------------------------------

def build(doc: Document) -> None:
    # Title
    title_p = doc.add_paragraph()
    r = title_p.add_run("Paper Trader — Project State Tracker")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = HEADING_BLUE
    r.font.name = BODY_FONT
    para(doc, "Canonical project-state doc. Built from the Claude Code audit "
              "(2026-05-11) plus strategic context — Polygon piggyback "
              "analysis, prioritization framing, and cleanup commit summary. "
              "Lives at docs/Project_State_Tracker.docx; update via Claude "
              "Code when state changes meaningfully.",
         italic=True, size=10)

    # ------------------------------------------------------------------
    # Executive Summary
    # ------------------------------------------------------------------
    h1(doc, "Executive Summary")
    para(doc, "Where the project actually stands as of May 11, 2026 "
              "(post-cleanup):")
    bullet_runs(doc, [
        ("Equity research framework (Mike) is production-stable.", True,
         False),
        (" Three promoted studies live on the cloud. NVDA/META "
         "concentration finding documented but not yet surfaced in "
         "dashboard exec summaries.", False, False)])
    bullet_runs(doc, [
        ("Crypto extension is paused.", True, False),
        (" 4-file stub committed, Chris pivoted to options. Crypto "
         "Phase 2 strategy code never started.", False, False)])
    bullet_runs(doc, [
        ("Options extension is the active build-out.", True, False),
        (" Chris shipped Sections 1-8 + 2.5 to main over 2 days. "
         "PR #15 (hang fix + concurrency refactor) is open and blocks "
         "the v1 production study from running.", False, False)])
    bullet_runs(doc, [
        ("Cloud dashboard renders equity content correctly but R2 is "
         "stale.", True, False),
        (" Last upload 2026-05-09; main has advanced 13 commits since "
         "(all options). Options content has never been uploaded to R2.",
         False, False)])
    bullet_runs(doc, [
        ("Mike's venv is missing dependencies.", True, False),
        (" truststore and pytest aren't installed — both are in "
         "requirements.txt. One-command fix.", False, False)])
    bullet_runs(doc, [
        ("SP1500 universe expansion remains stashed, awaiting Finnhub "
         "TOS clarity.", True, False),
        (" Polygon/Massive (Chris's data source) may make this question "
         "obsolete — see Section 3 for the piggyback math.",
         False, False)])

    para_runs(doc, [
        ("Honest read: ", True, False),
        ("The project is in a much better technical state than it was "
         "2 weeks ago. The mess is in coordination overhead — a stale "
         "R2, an open PR blocking a study, deferred decisions on "
         "infrastructure cleanup. The code itself is healthy. Three "
         "different research domains (equity, crypto, options) coexist "
         "without conflict thanks to Phase 1's asset-class refactor.",
         False, False)])

    h2(doc, "Status table at a glance")
    para(doc, "Red rows are URGENT (blocking work); yellow rows are warnings "
              "(needs attention but not blocking); green rows are working "
              "as designed.",
         italic=True, size=10)
    render_status_table(doc)

    # ------------------------------------------------------------------
    # Claude Session Onboarding
    # ------------------------------------------------------------------
    h1(doc, "Claude Session Onboarding")
    para(doc, "When opening a new Claude.ai or Claude Code session about "
              "this project, paste the prompt below into the first message. "
              "This gives Claude immediate context on project state, recent "
              "decisions, and what's in flight — so the session doesn't "
              "waste turns asking what the project is or rediscovering "
              "things this tracker already documents.")
    code_block(doc, ONBOARDING_PROMPT, bg_color=ONBOARDING_BG)
    para(doc, "After pasting, Claude will read this tracker and any "
              "referenced docs before answering your first question. Update "
              "this prompt's reference list if you add new authoritative "
              "docs to the repo.",
         italic=True, size=10)

    # ------------------------------------------------------------------
    # 1. Top 5 Findings
    # ------------------------------------------------------------------
    h1(doc, "1. Top 5 Most Important Findings")

    h2(doc, "Finding 1: PR #15 blocks the options v1 study")
    para(doc, "Branch chris/options-fix-v1-study-hang @ 7e0881c is 1 commit "
              "ahead of main. Real fix, not cosmetic. Root cause: engine.py "
              "was still passing Tradier's fetch_history to "
              "chain_reconstruction after PR #14 swapped the default to "
              "Polygon. Tradier returns null for expired OCCs, so "
              "reconstruct_chain looped on empty data.")
    para(doc, "Bundled fixes:")
    bullet(doc, "Separate RateLimiter per data source "
                "(Tradier 60/min vs Polygon 300/min)")
    bullet(doc, "Thread-safe lock on rate-limit deque "
                "(16 concurrent chain workers)")
    bullet(doc, "Capped header-driven sleep at 60s")
    bullet(doc, "Multi-day pre-fetch in _default_deps")
    bullet(doc, "Removed Polygon cache-write sanity gate that was rejecting "
                "legitimate sparse-expiry contracts")
    para(doc, "8 files, +526/-106 lines, new tests included. Impact: v1 "
              "production study cannot complete until this lands. No "
              "models/snapshots/options/pre_options_v1_<date>/ exists yet.")
    para_runs(doc, [
        ("Action: ", True, False),
        ("Mike reviews the diff, merges via squash on GitHub. "
         "~10 minutes.", False, False)])

    h2(doc, "Finding 2: R2 is stale — main is 13 commits ahead of last upload")
    para(doc, "R2 bucket snapshot_manifest.json says git_sha=bfee6dc, "
              "generated_at=2026-05-09T17:20:38Z. Main is at f3c3800. "
              "Functionally, equity content is still correct because none "
              "of those 13 commits touched equity files — they're all "
              "options work. But options content has never been uploaded "
              "to R2 at all (no options/ prefix exists in the bucket).")
    para_runs(doc, [("Action sequence:", True, False)])
    bullet(doc, "Merge PR #15")
    bullet(doc, "Run snapshot_for_cloud.py --asset-class equities (catches "
                "any equity changes — should be a near-no-op)")
    bullet(doc, "Run snapshot_for_cloud.py --asset-class options "
                "(first-ever options upload)")
    bullet(doc, "Hard-refresh the cloud dashboard, verify Options tab no "
                "longer shows placeholder")

    h2(doc, "Finding 3: Mike's venv is missing truststore and pytest")
    para(doc, "truststore==0.10.4 is in requirements.txt but not installed. "
              "Two test failures, both same root cause:")
    bullet(doc, "test_ssl_setup.py: collection error (ModuleNotFoundError: "
                "No module named 'truststore')")
    bullet(doc, "test_optuna_runner.py::test_smoke_study_constants_match_locked_config: "
                "fails for the same reason transitively via "
                "src/options/_ssl.py")
    para(doc, "After fixing: 374 tests pass, 1 skipped. Fix is one command:")
    code_block(doc, "venv\\Scripts\\pip install -r requirements.txt")
    para(doc, "Also flagged: pytest itself isn't in requirements.txt at "
              "all. Add it as a dev dep so future fresh venvs are "
              "self-bootstrapping.")

    h2(doc, "Finding 4: .env.example missing POLYGON_API_KEY — silent "
            "onboarding failure")
    para(doc, "src/options/polygon.py requires POLYGON_API_KEY and even "
              "cites Massive.com as the signup destination. But "
              ".env.example only documents Tradier env vars. A new "
              "contributor (or Mike on a fresh checkout) will hit a "
              "runtime error with no docs path forward.")
    para_runs(doc, [
        ("Easy fix: ", True, False),
        ("add to .env.example with a comment linking to "
         "https://polygon.io/signup.", False, False)])

    h2(doc, "Finding 5: Repo hygiene cleanup completed "
            "(2026-05-11, commit 5c39cc5)")
    para(doc, "This was the previous \"phantom CODEOWNERS paths + 3 orphan "
              "snapshots + 27 scratch files\" finding. Status as of "
              "cleanup commit 5c39cc5:")
    para_runs(doc, [("Done:", True, False)])
    bullet(doc, "35 ephemeral files deleted (27 debug scripts, 8 log files; "
                "all gitignored)")
    bullet(doc, "3 equity snapshots moved under "
                "models/snapshots/equities/ for consistency")
    bullet(doc, "3 tracked research scripts moved to scripts/research/")
    bullet(doc, "docs/sp1500_fetch_failures.txt → docs/diagnostics/")
    bullet(doc, "Dashboard_Viz_Implementation_Spec.docx → docs/specs/")
    bullet(doc, "Phantom /src/equities/ rule removed from CODEOWNERS")
    bullet(doc, "Stale remote branch origin/chris/options-phase-1-shared-edge "
                "deleted")
    bullet(doc, "Pre-existing gitignore bug fixed: !docs/*.docx was only "
                "1-depth; now !docs/**/*.docx covers all depths. This also "
                "auto-tracked docs/specs/Dashboard_Rebuild_Spec.docx which "
                "had been silently ignored for weeks.")
    para_runs(doc, [
        ("Net result: ", True, False),
        ("Repo root now has 5 files (requirements.txt, 3 SP1500 "
         "paused-state logs kept intentionally, v3_track2_runner.py "
         "which is referenced by dashboard). Clean state.",
         False, False)])

    # ------------------------------------------------------------------
    # 2. Recommended Action Sequence
    # ------------------------------------------------------------------
    h1(doc, "2. Recommended Action Sequence")
    para(doc, "Numbered in execution order. Items below are prioritized by "
              "what they unblock or risk if deferred.")

    h2(doc, "URGENT (do today)")
    bullet_runs(doc, [
        ("Merge PR #15.", True, False),
        (" Mike eyeballs the diff on GitHub, clicks Squash and Merge. "
         "~10 min.", False, False)])
    bullet_runs(doc, [
        ("Run pip install -r requirements.txt.", True, False),
        (" Fixes the 2 test failures and gets truststore. <1 min.",
         False, False)])
    bullet_runs(doc, [
        ("Sync R2 with --asset-class equities and --asset-class options.",
         True, False),
        (" Brings cloud dashboard current. ~5 min.", False, False)])
    bullet_runs(doc, [
        ("Visually verify the cloud dashboard.", True, False),
        (" Hard-refresh paper-trader-mteev.streamlit.app, click each tab. "
         "Confirm no 404 errors. ~3 min.", False, False)])

    h2(doc, "HIGH (this week)")
    bullet_runs(doc, [
        ("Add POLYGON_API_KEY to .env.example.", True, False),
        (" Trivial fix; eliminates silent onboarding failure. ~2 min.",
         False, False)])
    bullet_runs(doc, [
        ("Update Options_Extension_Decisions.md Status row.", True, False),
        (" Currently says \"Sections 1, 2, 3 merged\"; reality is 1-8 + "
         "2.5. Stale memos get ignored.", False, False)])
    bullet_runs(doc, [
        ("Surface NVDA/META finding in dashboard exec summaries.",
         True, False),
        (" Commit f46bd05 already exists locally with the research "
         "artifacts. Performance tab exec summary should say: ",
         False, False),
        ("\"Approximately 62% of this strategy's risk-adjusted edge "
         "comes from holdings in NVDA and META during the AI rally.\"",
         False, True),
        (" Switch headline alpha from arithmetic (+63.7pp/yr) to CAPM "
         "(+39.3pp/yr).", False, False)])
    bullet_runs(doc, [
        ("Decide on Polygon/Massive piggyback for stocks side.",
         True, False),
        (" See Section 3 for the analysis. This decision unblocks SP1500 "
         "and provides a long-term solution to yfinance unreliability.",
         False, False)])

    h2(doc, "MEDIUM (when convenient)")
    bullet_runs(doc, [
        ("Add a CI workflow (.github/workflows/test.yml).", True, False),
        (" Free with GitHub Actions; would have caught the "
         "truststore-vs-venv drift before it became a noticed bug.",
         False, False)])
    bullet_runs(doc, [
        ("Add pytest to requirements.txt", True, False),
        (" with a dev-only comment.", False, False)])

    h2(doc, "LOW (eventually)")
    bullet_runs(doc, [
        ("Resolve SP1500 stash@{0}.", True, False),
        (" Either pop and commit (if Finnhub clarity has arrived or you "
         "pivot to Polygon) or convert to a named branch.", False, False)])
    bullet_runs(doc, [
        ("Address optuna_runner.py:275 thread-safety race TODO.",
         True, False),
        (" Race is masked by workload heterogeneity but real.",
         False, False)])

    # ------------------------------------------------------------------
    # 3. Polygon Piggyback
    # ------------------------------------------------------------------
    h1(doc, "3. Massive (Polygon) Piggyback Analysis")
    para(doc, "This is the strategic question worth thinking about — whether "
              "to piggyback on Chris's Polygon subscription (rebranded as "
              "Massive since November 2025) for stocks-side data. The audit "
              "confirms Chris is already integrated with Polygon for the "
              "options work.")

    h2(doc, "What Chris's setup already does")
    bullet_runs(doc, [
        ("Authenticated client.", True, False),
        (" src/options/polygon.py is wrapped with proper rate limiting, "
         "retry-after handling, atomic cache writes.", False, False)])
    bullet_runs(doc, [
        ("Rate limiter pattern.", True, False),
        (" Separate RateLimiter instance per data source. The PR #15 "
         "refactor makes 16-thread concurrency safe. You could lift this "
         "pattern verbatim for equity work.", False, False)])
    bullet_runs(doc, [
        ("Cache layer.", True, False),
        (" Polygon responses cache to models/cache/options/polygon_cache/. "
         "Same pattern would apply to equity prices.", False, False)])

    h2(doc, "What Polygon offers for stocks (approximate, may have "
            "changed post-rebrand)")
    render_polygon_table(doc)
    para_runs(doc, [
        ("Caveat: ", True, False),
        ("Pricing above is from public sources and may have changed "
         "since the Massive rebrand. Chris should confirm what tier "
         "he's actually on before any piggyback decision.",
         False, True)])

    h2(doc, "What Polygon would solve")
    bullet_runs(doc, [
        ("yfinance silent-empty problem.", True, False),
        (" yfinance returns empty lists when throttled; Polygon returns "
         "proper errors with retry-after headers. The retry/sanity-gate "
         "code currently stashed for SP1500 becomes unnecessary because "
         "the failure mode it protects against doesn't exist on Polygon.",
         False, False)])
    bullet_runs(doc, [
        ("SP1500 universe expansion.", True, False),
        (" The pause reason was yfinance earnings dates returning 16% "
         "non-empty. Polygon's stocks plans include earnings dates as "
         "part of the API, not gated behind a separate product like "
         "Finnhub.", False, False)])
    bullet_runs(doc, [
        ("Historical data depth.", True, False),
        (" yfinance has 20+ years but data quality degrades for older "
         "periods, restated tickers get re-mapped silently, "
         "splits/dividends sometimes wrong. Polygon's historical data "
         "is exchange-sourced and stable.", False, False)])
    bullet_runs(doc, [
        ("Bulk download for cold-start runs.", True, False),
        (" Polygon offers daily flat-file S3 dumps. For a full-universe "
         "historical fetch, this is dramatically faster than per-ticker "
         "API calls.", False, False)])

    h2(doc, "Three piggyback paths, ranked")

    para_runs(doc, [
        ("Option A: ", True, False),
        ("Use Chris's existing key (single subscription, shared).",
         True, False),
        (" Mike uses Chris's POLYGON_API_KEY in Mike's .env. Both "
         "contributors hit the same Polygon account.", False, False)])
    bullet(doc, "Pros: No additional cost. Same data layer for both "
                "asset classes.")
    bullet(doc, "Cons: Mike's heavy stock fetches eat into Chris's "
                "options rate budget. If Mike runs an SP1500 fetch while "
                "Chris is running an options study, they throttle each "
                "other.")
    bullet(doc, "When this works: if Mike's equity work is mostly "
                "background research with predictable, slow-paced fetches.")

    para_runs(doc, [
        ("Option B: ", True, False),
        ("Mike subscribes to his own Polygon tier.", True, False),
        (" Two separate accounts. Each contributor pays their own "
         "subscription.", False, False)])
    bullet(doc, "Pros: Cleanest isolation. Predictable performance for both.")
    bullet(doc, "Cons: ~$29-79/mo for Mike's equity work depending on "
                "tier. Doubles the project's data costs.")
    bullet(doc, "When this works: if equity work is going to become "
                "serious / persistent.")

    para_runs(doc, [
        ("Option C: ", True, False),
        ("Hybrid — shared key for now, separate later.", True, False),
        (" Start with Option A. If quota fights become a problem, "
         "upgrade to Option B.", False, False)])
    bullet(doc, "Pros: Cheapest start. Real data informs the upgrade "
                "decision.")
    bullet(doc, "Cons: Need an early-warning mechanism (e.g., log all "
                "rate-limit hits) so the upgrade trigger doesn't get "
                "missed.")
    bullet(doc, "When this works: most likely the right answer for the "
                "current phase.")

    h2(doc, "Implementation cost")
    para(doc, "To pilot the piggyback approach for SP1500:")
    bullet(doc, "Write src/polygon_equities.py mirroring "
                "src/options/polygon.py structure — ~2-3 hours of "
                "Claude Code work")
    bullet(doc, "Pop the SP1500 stash; rewrite earnings_dates fetch to "
                "use Polygon instead of yfinance — ~1-2 hours")
    bullet(doc, "Run the SP1500 fetch against Polygon — if Polygon's "
                "free tier is enough for a one-time bulk download, this "
                "is $0; otherwise $29-79")
    bullet(doc, "Verify earnings data quality matches or beats what "
                "we'd get from yfinance")
    para(doc, "If the smoke test (~30 min) shows Polygon delivers usable "
              "earnings data, you've unblocked SP1500 expansion and "
              "removed the long-term yfinance reliability risk.")

    h2(doc, "Recommendation")
    para_runs(doc, [
        ("Take Option C.", True, False),
        (" Ask Chris what tier he's on and whether he's OK with you "
         "using his key. If yes, pilot the SP1500 fetch via Polygon as "
         "a single test case — that gives you concrete data on whether "
         "piggybacking is sustainable. If quota fights happen, switch "
         "to Option B. The total cost of being wrong is ~$50-100/mo "
         "— within budget for a research project this far along.",
         False, False)])

    # ------------------------------------------------------------------
    # 4. What this audit couldn't verify
    # ------------------------------------------------------------------
    h1(doc, "4. What This Audit Could Not Verify")
    para(doc, "Honesty section. These are things Claude Code flagged as "
              "unverifiable:")
    bullet(doc, "Whether the deployed Streamlit Cloud app currently "
                "renders without error in a browser. Manifest read "
                "suggests equity content is correct, but visual "
                "confirmation needed.")
    bullet(doc, "Whether models/cache/optuna_studies.db "
                "end-to-end-loads. File is 16 MB and looks intact, but "
                "the read path wasn't exercised.")
    bullet(doc, "Equity unit-test coverage. The 374 passing tests are "
                "mostly options-side. Equity correctness is asserted via "
                "reproducibility against the locked snapshot, not unit "
                "tests.")
    bullet(doc, "Whether Polygon API key in Mike's .env is valid. Audit "
                "didn't call the live endpoint.")
    bullet(doc, "Whether Chris has uncommitted local crypto work. Only "
                "main + chris/* remote branches were inspected.")

    # ------------------------------------------------------------------
    # 5. Strategic Read
    # ------------------------------------------------------------------
    h1(doc, "5. Strategic Read")
    para(doc, "Beyond the immediate action items, three things about where "
              "the project is now:")

    h2(doc, "The architecture is paying off")
    para(doc, "Three different research domains (equity, crypto, options) "
              "coexist in one codebase without conflict. The Phase 1 "
              "asset-class refactor is doing exactly what it was designed "
              "to do. Chris built an entire options module (21 files, 18 "
              "test files, 374 tests) without touching Mike's equity code, "
              "and Mike's equity work continued unaffected. That's the "
              "dividend on the upfront infrastructure investment.")

    h2(doc, "The bottleneck is coordination, not code")
    para(doc, "Everything in the URGENT and HIGH action lists is "
              "coordination work — review a PR, sync R2, update docs, "
              "install deps. No deep technical work is currently blocked. "
              "That's a much better state than \"we have a working "
              "strategy but the code is broken,\" which is where many "
              "quant projects get stuck.")

    h2(doc, "The next strategic decision is data infrastructure")
    para(doc, "The Polygon piggyback question is the most consequential "
              "decision pending. It affects:")
    bullet(doc, "SP1500 expansion (currently paused on yfinance reliability)")
    bullet(doc, "Future equity walk-forward validation")
    bullet(doc, "Cross-asset research consistency (using the same data "
                "vendor for both)")
    bullet(doc, "Eventually, live trading data quality")
    para(doc, "If you decide to go Polygon for stocks: that's a $0-79/mo "
              "decision that materially improves the project's long-term "
              "ceiling. If you decide to stay yfinance: SP1500 expansion "
              "needs a different solution and you're permanently exposed "
              "to yfinance's flakiness at scale.")

    h2(doc, "On Chris's options work")
    para(doc, "Chris shipped 13 commits in 2 days. That's a meaningful "
              "sprint and the audit shows the work is real — proper "
              "Greeks (Black-Scholes), engine, position model, "
              "BacktestConfig, FeeModel, Optuna runner, promotion gate, "
              "concentration analysis. The PR #15 hang fix is exactly "
              "the kind of bug that shows up when you actually try to run "
              "something end-to-end. Worth merging soon so he doesn't "
              "have to context-switch back to it later.")
    para(doc, "Once PR #15 merges and the v1 study runs, you'll have a "
              "third promotable workstream alongside the three equity "
              "studies. The dashboard will need parallel options tabs "
              "eventually (or at least a Performance tab that handles "
              "both asset classes); that's a future-work item Chris "
              "should think about as he gets close to v1 promotion.")

    # ------------------------------------------------------------------
    # Appendix A: Quick Reference
    # ------------------------------------------------------------------
    h1(doc, "Appendix A: Quick Reference")

    h2(doc, "Action sequence — copy-paste")
    code_block(doc,
               "# 1. Install missing dependencies\n"
               "cd C:\\Users\\mteev\\Documents\\Claude-Investing\\paper_trader\n"
               "venv\\Scripts\\pip install -r requirements.txt\n"
               "\n"
               "# 2. After PR #15 merges, pull main\n"
               "git checkout main\n"
               "git pull origin main\n"
               "\n"
               "# 3. Sync R2 for both asset classes\n"
               "venv\\Scripts\\python.exe src\\snapshot_for_cloud.py --asset-class equities --dry-run\n"
               "venv\\Scripts\\python.exe src\\snapshot_for_cloud.py --asset-class equities\n"
               "\n"
               "venv\\Scripts\\python.exe src\\snapshot_for_cloud.py --asset-class options --dry-run\n"
               "venv\\Scripts\\python.exe src\\snapshot_for_cloud.py --asset-class options\n"
               "\n"
               "# 4. Hard-refresh https://paper-trader-mteev.streamlit.app\n"
               "#    Click each tab. Verify no errors. Verify Options tab shows\n"
               "#    real content instead of placeholder.\n"
               "\n"
               "# 5. Run pytest to confirm all tests pass\n"
               "venv\\Scripts\\pytest -v")

    h2(doc, "Key file paths")
    code_block(doc,
               "# Production data\n"
               "models/snapshots/equities/pre_v2_20260505/\n"
               "\n"
               "# Promoted studies (R2 reads these)\n"
               "models/cache/dashboard_results/best_*/\n"
               "\n"
               "# Locked configs\n"
               "models/cache/dashboard_results/best_regime_dependent_v1_20260505_2240_325/\n"
               "models/cache/dashboard_results/best_15_position_study_v1_842/\n"
               "models/cache/dashboard_results/best_continuous_sizing_v1_1852/\n"
               "\n"
               "# Documentation\n"
               "docs/Comprehensive_User_Guide.docx\n"
               "docs/Crypto_Extension_Decisions.md\n"
               "docs/Options_Extension_Decisions.md\n"
               "docs/future_work.md\n"
               "docs/Project_State_Tracker.docx  (this doc, committed)\n"
               "\n"
               "# Specs for next Claude Code sessions\n"
               "docs/specs/Sizing_Decisions_Viz_Prompt.md\n"
               "docs/specs/Dashboard_Rebuild_Spec.docx\n"
               "docs/specs/Dashboard_Viz_Implementation_Spec.docx\n"
               "\n"
               "# Research scripts (one-off, history preserved)\n"
               "scripts/research/backfill_benchmark_parquets.py\n"
               "scripts/research/research_blacklist_325.py\n"
               "scripts/research/sanity_test_continuous_sizing.py\n"
               "\n"
               "# Audit / state tooling\n"
               "scripts/build_project_state_audit.py    (regenerator for state audits)\n"
               "scripts/build_project_state_tracker.py  (regenerator for this tracker)")

    h2(doc, "Branch states")
    code_block(doc,
               "# Active branches\n"
               "main                                       # latest commit\n"
               "chris/options-fix-v1-study-hang            # 7e0881c (1 ahead, OPEN PR #15)\n"
               "\n"
               "# Stashed work\n"
               "stash@{0}                                  # SP1500 retry/sanity-gate\n"
               "                                           # awaits Finnhub TOS or Polygon decision")

    # ------------------------------------------------------------------
    # Appendix B: Document Lineage
    # ------------------------------------------------------------------
    h1(doc, "Appendix B: Document Lineage")
    para(doc, "This is the canonical project state doc. It supersedes "
              "prior versions and is committed to the repo for the first "
              "time.")
    render_lineage_table(doc)
    para_runs(doc, [
        ("Older docs remain valid as historical reference. ", False, False),
        ("This doc is the most current source of truth for current state "
         "and immediate next steps.", True, False),
        (" For deep technical reference, fall back to "
         "Comprehensive_User_Guide.docx. For audit-level detail Claude "
         "Code produced, fall back to Project_State_Audit.docx.",
         False, False)])
    para(doc, "This doc lives at docs/Project_State_Tracker.docx in the "
              "repo. Update via Claude Code when state changes "
              "meaningfully — e.g., after PR #15 merges, after a major "
              "Polygon decision, after the next big sprint.")
    para_runs(doc, [
        ("Maintenance: ", True, False),
        ("when adding new authoritative docs to docs/, update the "
         "\"Claude Session Onboarding\" prompt above to reference them. "
         "When contributor roles or conventions change, update the "
         "prompt's conventions section. Specifically, review the embedded "
         "prompt whenever:", False, False)])
    bullet(doc, "Major docs are added to or removed from the reading list")
    bullet(doc, "Project conventions change (data layout, branching, "
                "deployment targets)")
    bullet(doc, "Contributor roles shift (ownership, handoffs, "
                "namespace boundaries)")

    # ------------------------------------------------------------------
    # Appendix C: Cleanup Session Summary
    # ------------------------------------------------------------------
    h1(doc, "Appendix C: Cleanup Session Summary (2026-05-11)")
    para_runs(doc, [
        ("Commit: ", True, False),
        ("5c39cc5", False, False),
        (" — \"chore: cleanup per project state audit\"", False, False)])

    para_runs(doc, [("Deleted (35 files, gitignored):", True, False)])
    bullet(doc, "27 debug/inspection Python scripts at repo root "
                "(inspect_*.py, check_*.py, track_*.py, study_progress.py, "
                "progress_continuous.py, touch_snapshot.py, "
                "rescue_save_continuous.py, best_continuous.py, "
                "spot_check.py, compare_studies.py, extract_track2*.py, "
                "make_src_zip.py)")
    bullet(doc, "8 log files (15_position_study_v1*.log, "
                "continuous_sizing_v1*.log, v3_track1_*.log)")
    bullet(doc, "src.zip (46 MB)")
    bullet(doc, "_spec_extracted.txt")
    bullet(doc, "Documents/ folder (3 early-draft files) and empty "
                "data/ directory")

    para_runs(doc, [("Moved:", True, False)])
    bullet(doc, "3 snapshots (macro_signal_investigation_*, "
                "parallelism_diagnostic_*, post_macro_fix_*) → "
                "models/snapshots/equities/")
    bullet(doc, "3 tracked research scripts → scripts/research/")
    bullet(doc, "docs/sp1500_fetch_failures.txt → docs/diagnostics/")
    bullet(doc, "Dashboard_Viz_Implementation_Spec.docx → docs/specs/")

    para_runs(doc, [("Kept (intentional):", True, False)])
    bullet(doc, "3 SP1500 paused-state logs (sp1500_fetch.log, "
                "sp1500_fetch_v2.log, sp1500_force_refresh.log) — "
                "document the paused state for when SP1500 resumes")
    bullet(doc, "v3_track2_runner.py at root — referenced by "
                "src/dashboard_app.py")
    bullet(doc, "Chris's 6 options scripts in scripts/")

    para_runs(doc, [("Newly tracked (fixed by gitignore fix):", True, False)])
    bullet(doc, "docs/specs/Dashboard_Rebuild_Spec.docx — had been "
                "silently ignored due to !docs/*.docx depth bug")
    bullet(doc, "docs/specs/Dashboard_Viz_Implementation_Spec.docx")
    bullet(doc, "scripts/build_project_state_audit.py — for future audit "
                "refreshes")
    bullet(doc, "scripts/build_project_state_tracker.py — for future "
                "tracker refreshes (this doc's generator)")

    para_runs(doc, [("Modified:", True, False)])
    bullet(doc, ".gitignore: fixed !docs/*.docx to !docs/**/*.docx "
                "(was a depth-1 bug); added sp1500_*.log pattern")
    bullet(doc, ".github/CODEOWNERS: removed phantom /src/equities/ rule")

    para_runs(doc, [("Branches deleted:", True, False)])
    bullet(doc, "origin/chris/options-phase-1-shared-edge "
                "(already squash-merged via PR #3)")

    para_runs(doc, [
        ("Net result: ", True, False),
        ("Repo root now has 5 files (requirements.txt, 3 SP1500 paused "
         "logs, v3_track2_runner.py). Clean state.", False, False)])

    para_runs(doc, [
        ("Important secondary finding: ", True, False),
        ("The !docs/*.docx gitignore rule was buggy and had been "
         "silently dropping any docx file under docs/specs/ for weeks. "
         "The fix means committed documentation now reliably lands in "
         "the repo.", False, False)])


def main() -> None:
    doc = Document()
    configure_document(doc)
    build(doc)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}  ({OUT_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
