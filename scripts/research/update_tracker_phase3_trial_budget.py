"""Small follow-up tracker update — pin the Phase 3 trial budget specifics.

After Mike approved Option 2 (XGB 200 + ENet 100), update the Phase 3
section's status paragraph so partners checking in mid-run know exactly
what's running and when to expect completion.
"""
from __future__ import annotations

from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TRACKER = ROOT / "docs" / "Project_State_Tracker.docx"


def main() -> int:
    d = Document(TRACKER)

    # Find the section heading
    section_heading_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("Larger Universe v1 study"):
            section_heading_idx = i
            break
    if section_heading_idx is None:
        raise SystemExit("section heading not found")

    # Status paragraph is the first Normal paragraph after the heading
    status_para = None
    for i in range(section_heading_idx + 1, len(d.paragraphs)):
        if d.paragraphs[i].style.name == "Normal":
            status_para = d.paragraphs[i]
            break
    if status_para is None:
        raise SystemExit("status paragraph not found")

    # Compose new status text
    now_utc = datetime.now(timezone.utc).replace(microsecond=0)
    eta = now_utc + timedelta(hours=7, minutes=30)
    new_status = (
        f"Status: Phase 3 (Optuna hyperparameter tuning) authorized 2026-05-11 evening "
        f"with revised spec. Mike chose Option 3 from the horizon diagnostic — monthly "
        f"rebalance + monthly label horizon — putting the modeling cadence where the "
        f"data shows signal lives. Trial budget: 200 XGBoost trials + 100 ElasticNet "
        f"trials (complexity-asymmetric; ENet has 2 hyperparameters and TPE typically "
        f"plateaus by trial 50-80 on that search space, vs XGBoost's 9). Sequential "
        f"execution (XGB first, then ENet) to avoid CPU contention. Estimated wall-clock "
        f"~7.5 hours; backgrounded {now_utc.isoformat()}, expected completion "
        f"approximately {eta.isoformat()}. Fixed TPE sampler seed (42) for reproducibility. "
        f"Per-trial timing and convergence checkpoints logged every 25 trials at "
        f"models/studies/larger_universe_v1/phase3_progress.log."
    )

    # Replace status paragraph contents
    for r in status_para._element.findall(qn("w:r")):
        status_para._element.remove(r)
    status_para.add_run(new_status)

    d.save(TRACKER)
    print(f"updated {TRACKER}")
    print(f"backgrounded_at: {now_utc.isoformat()}")
    print(f"expected_completion: {eta.isoformat()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
