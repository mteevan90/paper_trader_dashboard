"""Update Project_State_Tracker.docx to reflect Phase 3 authorization.

Replaces the "paused at Phase 2→3 gate" status (added last session) with
"Phase 3 running with revised spec (21d label, monthly rebalance)".

Strategy: rather than re-edit the existing paragraphs, we update the
heading text and rewrite the status paragraph; the rest of the section
(what's complete, three options, partners note) is replaced with a
shorter "current state" + a pointer to the session log for granular detail.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TRACKER = ROOT / "docs" / "Project_State_Tracker.docx"


def main() -> int:
    d = Document(TRACKER)

    # Find the section heading "Larger Universe v1 study — paused at Phase 2→3 gate..."
    section_heading_idx = None
    next_h1_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("Larger Universe v1 study"):
            section_heading_idx = i
        elif section_heading_idx is not None and p.style.name == "Heading 1":
            next_h1_idx = i
            break

    if section_heading_idx is None or next_h1_idx is None:
        raise SystemExit("could not locate the Larger Universe v1 section in tracker")

    # Update the heading text
    heading = d.paragraphs[section_heading_idx]
    for r in heading._element.findall(qn("w:r")):
        heading._element.remove(r)
    heading.add_run("Larger Universe v1 study — Phase 3 running with revised spec (2026-05-11 evening)")

    # Replace section body paragraphs (everything between section_heading_idx
    # and next_h1_idx) by deleting them, then inserting fresh content.
    # Walk backward so indices don't shift.
    body_para_elems = []
    for i in range(section_heading_idx + 1, next_h1_idx):
        body_para_elems.append(d.paragraphs[i]._element)
    for el in body_para_elems:
        el.getparent().remove(el)

    # Insert new content. After each insert, "after" becomes the new para
    # so subsequent inserts land in document order.
    from docx.text.paragraph import Paragraph

    def insert_after(ref_paragraph, text, style):
        ref_elem = ref_paragraph._element
        new_p = deepcopy(ref_elem)
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        ref_elem.addnext(new_p)
        par = Paragraph(new_p, ref_paragraph._parent)
        par.style = par.part.document.styles[style]
        par.add_run(text)
        return par

    items = [
        ("Normal",
         "Status: Phase 3 (Optuna hyperparameter tuning) authorized 2026-05-11 evening with revised spec. Mike chose Option 3 from the horizon diagnostic — monthly rebalance + monthly label horizon — putting the modeling cadence where the data shows signal lives. Expected wall-clock ~10-12 hours; surfaced for trial-count confirmation before backgrounding."),
        ("Heading 2", "Spec revision summary"),
        ("List Bullet",
         "Label horizon: 5 trading days → 21 trading days (~monthly forward return)"),
        ("List Bullet",
         "Rebalance cadence: weekly (Friday close) → monthly (last trading day)"),
        ("List Bullet",
         "Rebalance threshold: 1.5pp → REMOVED (rebalance fully each month at the lower turnover frequency; FeeModel accrues costs on actual position changes)"),
        ("List Bullet",
         "CV embargo: 5 days → 21 days (= new label horizon, per finance-ML convention)"),
        ("Normal",
         "All other parameters unchanged: train 2017-05-12 to 2023-05-11, test 2023-05-12 to 2025-12-31, OOS holdout 2026-01-01 onward; 5-fold expanding-window CV; XGBoost primary + ElasticNet sanity; 7.5%/30% position/sector caps; four benchmarks; 0.05% flat FeeModel; long-only; score-weighted continuous sizing."),
        ("Heading 2", "Why the revision"),
        ("Normal",
         "Phase 2 horizon diagnostic (variants A and B, see docs/diagnostics/larger_universe_v1_horizon_diagnostic.md) showed cross-sectional alpha lives at monthly horizons. At 5-day horizon: XGBoost collapsed to constant-within-date predictions (learned macro features, ignored ticker-level features); ElasticNet mean cross-sectional IC 0.020 with high per-date noise. At 21-day horizon: XGBoost no longer degenerates, mean IC 0.019 with all 5 folds covered; ElasticNet 0.031 with 4 of 5 folds positive. The horizon shift was a bigger lift than adding the full fundamentals/sector feature set at 5d. Consistent with the academic factor-research literature."),
        ("Heading 2", "For Mike's partners checking in"),
        ("Normal",
         "Phase 3 is the multi-hour tuning run; it can be safely interrupted by reading the partial logs at models/studies/larger_universe_v1/phase3_progress.log. Best parameters and intermediate study state persist every 10 trials at models/studies/larger_universe_v1/{xgboost,elasticnet}_study.json. The granular session-by-session activity log is at docs/sessions/larger_universe_v1/session_log.md. Locked spec at docs/studies/larger_universe_v1/spec.md. After Phase 3 completes, Mike reviews before authorizing Phase 4 (portfolio construction + backtest)."),
    ]
    ref = d.paragraphs[section_heading_idx]
    for style, text in items:
        ref = insert_after(ref, text, style)

    d.save(TRACKER)
    print(f"updated {TRACKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
