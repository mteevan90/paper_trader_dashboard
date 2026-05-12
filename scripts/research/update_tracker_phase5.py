"""Update Project_State_Tracker.docx to reflect full Larger Universe v1 completion.

Replaces the existing "Larger Universe v1 study — Phase 3 running with revised spec"
section (paragraphs 17-28) with a new section documenting the completed study.
Leaves everything else in the tracker untouched.
"""
from __future__ import annotations

import shutil
from pathlib import Path

import docx
from docx.oxml.ns import qn

TRACKER = Path(__file__).resolve().parents[2] / "docs" / "Project_State_Tracker.docx"
BACKUP = TRACKER.with_suffix(".docx.bak_phase5_update")


def _find_section_bounds(doc: docx.Document) -> tuple[int, int]:
    """Return (start_idx, end_idx) of the Larger Universe v1 Phase 3 section."""
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        if (p.style.name == "Heading 1"
                and "Larger Universe v1 study" in p.text):
            start = i
        elif start is not None and p.style.name == "Heading 1":
            end = i
            break
    if start is None:
        raise SystemExit("Did not find the Larger Universe v1 section.")
    if end is None:
        raise SystemExit("Did not find the section that follows.")
    return start, end


def _delete_paragraphs(doc: docx.Document, start: int, end: int) -> None:
    """Delete paragraphs[start:end] from the document body."""
    paragraphs = doc.paragraphs[start:end]
    for p in paragraphs:
        p._element.getparent().remove(p._element)


def _insert_paragraph_before(anchor_para, text: str, style: str):
    """Insert a new paragraph with `text` and `style` before `anchor_para`."""
    new_p = anchor_para.insert_paragraph_before(text=text, style=style)
    return new_p


def main() -> None:
    shutil.copy2(TRACKER, BACKUP)
    print(f"Backup written to {BACKUP}")

    doc = docx.Document(str(TRACKER))
    start, end = _find_section_bounds(doc)
    print(f"Replacing paragraphs [{start}, {end}) — "
          f"{end - start} paragraphs.")

    # Capture anchor (the heading that follows our section)
    anchor = doc.paragraphs[end]

    # Delete the old section
    _delete_paragraphs(doc, start, end)

    # Refresh paragraph list after deletion
    doc_after_delete = doc

    # Find the anchor again by text
    anchor_text = anchor.text
    anchor = None
    for p in doc_after_delete.paragraphs:
        if p.text == anchor_text and p.style.name == "Heading 1":
            anchor = p
            break
    assert anchor is not None, "Anchor lost after delete"

    # Build new section content
    new_content = [
        ("Heading 1",
         "Larger Universe v1 study — COMPLETE (not promoted) — 2026-05-12"),
        ("Normal",
         "Status: Phase 5 complete on branch feat/larger-universe-v1-study "
         "(pre-merge, awaiting Mike's review). Branch contains the full "
         "five-phase study build plus an architectural memo, a dashboard-"
         "contract update, a Phase 4.5 dashboard implementation, six "
         "static figures, and this tracker update — all bundled as one "
         "coherent unit. Merge decision happens after review."),

        ("Heading 2", "Headline disposition"),
        ("List Bullet",
         "Universe expansion succeeded: 2,122-ticker SP1500+delisted "
         "snapshot with documented survivorship-bias mitigation. "
         "Infrastructure is reusable for future studies."),
        ("List Bullet",
         "Neither model met all hard success criteria. XGBoost (primary) "
         "earned +3.5pp excess CAGR vs SPY on the test window but failed "
         "the 1.5×-SPY-drawdown constraint (−33.5% vs −28.5% threshold) "
         "and the 25%-single-ticker constraint (MXL at 33.9%). ElasticNet "
         "(sanity) showed +21.2pp excess CAGR but 88% of its alpha came "
         "from a single ticker (DBD)."),
        ("List Bullet",
         "Per spec: not promoted. The strategy beats SPY on average but "
         "fails risk-adjusted promotion criteria. Documented for the "
         "record; not deployed."),
        ("List Bullet",
         "Two methodology findings worth retaining: (1) full-cross-section "
         "Spearman IC was the wrong CV objective for top-N portfolio "
         "strategies — XGBoost CV winner had +0.028 in-fold IC but −0.009 "
         "held-out full-IC with +0.048 top-quintile IC; (2) feature/model/"
         "construction interaction drove the DBD concentration "
         "deterministically (trend features + linear model + light "
         "regularization + rank-top-N)."),
        ("List Bullet",
         "First contract-conformant study under dashboard_contract_v1. "
         "Future studies inherit the dashboard infrastructure rather than "
         "each retrofitting it."),

        ("Heading 2", "What was built (Phase 5 close session, 2026-05-12)"),
        ("List Bullet",
         "docs/architecture/ml_study_cv_objectives_v1.md — architectural "
         "memo distilling the CV-objective finding. APPROVED 2026-05-12. "
         "Establishes a dual-reporting requirement so future studies can "
         "validate or revise the top-quintile-IC recommendation (n=1 today)."),
        ("List Bullet",
         "docs/architecture/dashboard_contract_v1.md — updated with the "
         "objective.training_cv field requirement under meta.json.objective. "
         "Controlled vocabulary added (top_quintile_spearman_ic recommended)."),
        ("List Bullet",
         "docs/studies/larger_universe_v1/results.md — full study writeup "
         "with executive summary, methodology (Phase 1-5), six embedded "
         "figures, the DBD case study, success-criteria pass/fail table, "
         "forward-looking findings, future-work candidates (no endorsement), "
         "methodology limitations."),
        ("List Bullet",
         "docs/studies/larger_universe_v1/figures/*.png — six static "
         "figures: equity_curves, year_by_year, decile_returns, "
         "alpha_attribution, walk_forward, ic_decomposition. "
         "Generated via scripts/research/phase5_generate_figures.py."),
        ("List Bullet",
         "src/dashboard_app.py — Phase 4.5 dashboard. Sidebar now offers a "
         "Legacy / Contract-conformant radio (Legacy is default). The "
         "Contract-conformant branch renders 8 tabs from contract_v1/ "
         "artifacts: Overview, Performance, Holdings, Trades, Alpha "
         "Attribution, Diagnostics, Walk-forward, Tuning. Legacy tabs "
         "unchanged."),
        ("List Bullet",
         "models/studies/larger_universe_v1/contract_v1/ — 14 artifacts "
         "(meta.json + 12 parquet/json files) following the dashboard "
         "contract. Phase 4 and Phase 5 analytics outputs."),

        ("Heading 2", "Pointers"),
        ("List Bullet",
         "Writeup: docs/studies/larger_universe_v1/results.md"),
        ("List Bullet",
         "Session log: docs/sessions/larger_universe_v1/session_log.md"),
        ("List Bullet",
         "Architectural memo: docs/architecture/ml_study_cv_objectives_v1.md"),
        ("List Bullet",
         "Dashboard contract: docs/architecture/dashboard_contract_v1.md"),
        ("List Bullet",
         "Spec: docs/studies/larger_universe_v1/spec.md (Phase 4 spec at "
         "phase4_spec.md)"),
        ("List Bullet",
         "Contract artifacts: models/studies/larger_universe_v1/contract_v1/"),
        ("List Bullet",
         "Dashboard route: /studies/larger_universe_v1 via the "
         "Contract-conformant sidebar branch in src/dashboard_app.py"),

        ("Heading 2", "What's next (post-review)"),
        ("List Bullet",
         "Mike reviews the coherent unit on feat/larger-universe-v1-study. "
         "Approve → merge to main via PR (Mike handles). Request revisions → "
         "Claude addresses specific revisions before merge."),
        ("List Bullet",
         "v2 study design conversation deferred until after merge. Not "
         "queued or specced. Future-work section in the writeup lists "
         "candidate directions without endorsement."),
        ("List Bullet",
         "Any future contract-conformant study inherits the Phase 4.5 "
         "dashboard tabs by writing artifacts to its own contract_v1/ "
         "folder per the contract spec."),

        ("Heading 2", "For Mike's partners checking in"),
        ("Normal",
         "Phase 5 is done. Branch sits pre-merge for Mike's review. There "
         "is no in-flight long-running compute; the work to look at is "
         "documentation (writeup, memo, session log), dashboard code, and "
         "the contract artifacts. The promotion criteria came in as a clean "
         "fail on two of three hard criteria for both models, which means "
         "the study contributes methodology learnings rather than a new "
         "promoted strategy — that outcome is documented honestly and not "
         "spun."),
    ]

    for style, text in new_content:
        _insert_paragraph_before(anchor, text, style)

    # Save
    doc.save(str(TRACKER))
    print(f"Updated tracker saved to {TRACKER}.")


if __name__ == "__main__":
    main()
