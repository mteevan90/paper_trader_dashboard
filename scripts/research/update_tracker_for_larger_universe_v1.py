"""Update Project_State_Tracker.docx to reflect Larger Universe v1.

Two changes:

1. Insert an OBSOLETE marker at the top of Section 3 (Massive/Polygon
   piggyback analysis). Finnhub Basic is the chosen vendor for stocks-side
   data; Chris's Polygon subscription is options-only with 2-year history
   and cannot be repurposed.

2. Insert a new sub-section into Appendix A: Quick Reference that lists
   the two equity snapshots and their roles:
     - pre_v2_20260505: legacy baseline (locked, used by promoted studies)
     - larger_universe_v1_20260511: best-effort survivorship-bias
       mitigation, fresh study material
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TRACKER = ROOT / "docs" / "Project_State_Tracker.docx"


def insert_paragraph_after(reference_paragraph, text, style=None):
    """Insert a paragraph after the given paragraph, return the new paragraph."""
    # Work on the underlying XML element
    ref_elem = reference_paragraph._element
    new_p = deepcopy(ref_elem)
    # Clear runs from the cloned XML
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    ref_elem.addnext(new_p)
    # Wrap so we can set text + style
    from docx.text.paragraph import Paragraph
    par = Paragraph(new_p, reference_paragraph._parent)
    if style is not None:
        par.style = par.part.document.styles[style]
    par.add_run(text)
    return par


def main() -> int:
    d = Document(TRACKER)

    # Find Section 3 heading
    section_3_idx = None
    appendix_a_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("3. Massive"):
            section_3_idx = i
        if p.style.name == "Heading 1" and p.text.strip() == "Appendix A: Quick Reference":
            appendix_a_idx = i

    if section_3_idx is None or appendix_a_idx is None:
        raise SystemExit("could not locate Section 3 and/or Appendix A in tracker")

    # 1. Mark Section 3 obsolete by inserting a paragraph after the heading.
    s3_heading = d.paragraphs[section_3_idx]
    insert_paragraph_after(
        s3_heading,
        (
            "OBSOLETE (2026-05-11): This analysis is superseded. Chris's Polygon "
            "subscription is options-only with a 2-year history limit and cannot be "
            "repurposed for stocks-side equity data. Mike instead purchased Finnhub "
            "Basic ($49.99/mo, 150 calls/min for candle, 60 calls/min for /stock/metric, "
            "10y daily OHLC, survivorship-bias-mitigated). The Larger Universe v1 "
            "snapshot at models/snapshots/equities/larger_universe_v1_20260511/ is the "
            "result. See docs/diagnostics/larger_universe_v1_capabilities.md and "
            "docs/diagnostics/larger_universe_v1_snapshot_summary.md for details. The "
            "remainder of this section is preserved as historical record of the "
            "decision context."
        ),
        style="Normal",
    )

    # 2. Append new Appendix A sub-section listing both equity snapshots.
    # Find the LAST paragraph inside Appendix A (just before Appendix B).
    appendix_b_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("Appendix B"):
            appendix_b_idx = i
            break
    if appendix_b_idx is None:
        raise SystemExit("could not locate Appendix B")

    insert_after = d.paragraphs[appendix_b_idx - 1]
    # Insert in reverse order so document flow ends up correct
    items_in_order = [
        ("Heading 2", "Equity snapshots — Larger Universe v1 added 2026-05-11"),
        ("Normal",
         "models/snapshots/equities/pre_v2_20260505/  — legacy baseline (locked). "
         "491 ticker universe (SP500 + NDX overlap), 2018-01-02 onward, "
         "survivorship-biased. Used by the three promoted studies (#325, #842, #1852). "
         "Do not modify."),
        ("Normal",
         "models/snapshots/equities/larger_universe_v1_20260511/  — Larger Universe "
         "v1 baseline (best-effort survivorship-bias mitigation, fresh study material). "
         "1,963 tickers with 10y daily prices, 1,919 with Finnhub fundamentals, "
         "macro_signals carried forward. SP500 + SP400 + SP600 actives plus "
         "last-decade delisted. No earnings_dates (Finnhub Basic forward-only; "
         "yfinance scale-fetch hit 23.6% coverage and the stash sanity gate "
         "fired — feature dropped from v1 per the <85% rule). See the snapshot "
         "summary doc for the full coverage breakdown and residual-bias "
         "characterization."),
    ]
    # Insert each in turn; the previous one becomes the new reference
    ref = insert_after
    for style, text in items_in_order:
        ref = insert_paragraph_after(ref, text, style=style)

    d.save(TRACKER)
    print(f"updated {TRACKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
