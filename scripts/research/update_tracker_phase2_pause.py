"""Update Project_State_Tracker.docx for the Larger Universe v1 Phase 2 pause.

Per the standing rule: tracker gets updated before any long unattended run.
Phase 3 is deferred tonight pending diagnostic review, so the tracker
should reflect "paused at Phase 2→3 gate" rather than "in execution".

Changes:
1. Add a "Larger Universe v1 study (paused at Phase 2→3 gate)" section
   near the top, between Executive Summary and Findings, since this is
   the most-current state of the project.
2. Add a partners-note at the very top describing where to look for
   current state.
3. Add a session-log pointer in Appendix A.

Data files are unchanged; this is documentation only.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TRACKER = ROOT / "docs" / "Project_State_Tracker.docx"


def insert_paragraph_after(reference_paragraph, text, style=None):
    """Insert a paragraph after the given paragraph, return the new paragraph."""
    ref_elem = reference_paragraph._element
    new_p = deepcopy(ref_elem)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    ref_elem.addnext(new_p)
    from docx.text.paragraph import Paragraph
    par = Paragraph(new_p, reference_paragraph._parent)
    if style is not None:
        par.style = par.part.document.styles[style]
    par.add_run(text)
    return par


def main() -> int:
    d = Document(TRACKER)

    # Locate landmarks
    exec_summary_idx = None  # the "Executive Summary" Heading 1
    findings_h1_idx = None   # the "1. Top 5 Most Important Findings" Heading 1
    appendix_a_idx = None
    appendix_b_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.style.name == "Heading 1":
            t = p.text.strip()
            if t == "Executive Summary":
                exec_summary_idx = i
            elif t.startswith("1. Top 5"):
                findings_h1_idx = i
            elif t == "Appendix A: Quick Reference":
                appendix_a_idx = i
            elif t.startswith("Appendix B"):
                appendix_b_idx = i

    if findings_h1_idx is None:
        raise SystemExit("could not locate findings heading")
    if appendix_a_idx is None or appendix_b_idx is None:
        raise SystemExit("could not locate appendix landmarks")

    # ----- Insertion 1: Larger Universe v1 status section BEFORE the
    # Findings heading (so it appears between Executive Summary content
    # and the historical Findings 1-5). We insert by placing items AFTER
    # the paragraph immediately preceding findings_h1.
    findings_h1 = d.paragraphs[findings_h1_idx]
    ref_for_insert = d.paragraphs[findings_h1_idx - 1]

    # Items inserted in order — they'll appear in reverse-insert order, so
    # we have to walk the reference forward as we insert.
    items = [
        ("Heading 1", "Larger Universe v1 study — paused at Phase 2→3 gate (2026-05-11)"),
        ("Normal",
         "Status: paused at Phase 2→3 gate pending diagnostic re-smoke results. The cross-sectional IC bug fix surfaced near-zero stock-ranking signal on the price+macro feature subset at 5-day horizon; horizon diagnostic shows signal exists at 21-day horizon for the full feature set. Phase 3 (6-7 hour Optuna run) deferred to Mike's morning review."),
        ("Heading 2", "What's complete"),
        ("List Bullet",
         "Larger Universe v1 data snapshot at models/snapshots/equities/larger_universe_v1_20260511/ — 1,963 tickers, 10y daily prices, 1,919 with Finnhub fundamentals, survivorship-bias-mitigated via Wikipedia component-change scrape."),
        ("List Bullet",
         "Feature matrix at models/features/larger_universe_v1/features.parquet — 4.35M rows × 40 cols including point-in-time fundamentals (45d reporting lag), extended macro signals (10 FRED series), sector + index membership, derived log market cap, rolling-36-month beta."),
        ("List Bullet",
         "Training pipelines (XGBoost + ElasticNet) with 5-fold expanding-window CV and 5d embargo. CV design doc at docs/diagnostics/larger_universe_v1_cv_design.md."),
        ("List Bullet",
         "Cross-sectional IC bug found and fixed pre-Phase-3. Panel-wise IC of 0.085 was substantially market-timing rather than stock-ranking. Corrected metric correctly attributes zero credit to stock-ranking on the price+macro subset."),
        ("Heading 2", "Headline diagnostic finding"),
        ("Normal",
         "Horizon is the binding constraint, not feature set. At 5-day horizon, XGBoost collapses to constant-within-date predictions (learns macro features, ignores ticker-level features). At 21-day horizon, both models produce meaningful cross-sectional IC across all folds: XGBoost mean 0.019, ElasticNet mean 0.031, with positive contributions from 4 of 5 folds. Fold 3 (val 2021-05 to 2022-05) is hostile for both models — the 2022 bear-market regime reversal. See docs/diagnostics/larger_universe_v1_horizon_diagnostic.md for the full per-fold breakdown."),
        ("Heading 2", "Decision pending — three options"),
        ("List Bullet",
         "Option 1 (recommended): change label horizon to 21d, keep weekly rebalance. Model predicts 21-day forward return; portfolio rebalances weekly using the most recent prediction. Each prediction informs ~4 rebalances before becoming stale. Uses signal where it lives, preserves spec's weekly cadence."),
        ("List Bullet",
         "Option 2: keep 5d label, run Phase 3 anyway, accept low IC. Phase 3 burns 6-7h to tune on a near-zero signal landscape. Expected to underperform Option 1."),
        ("List Bullet",
         "Option 3: pivot to monthly rebalance + monthly label. Cleanest factor-research design but violates the spec's weekly cadence."),
        ("Heading 2", "For Mike's partners checking in while Mike is away"),
        ("Normal",
         "Current state: paused, by design. The Phase 2→3 gate caught a measurement issue (cross-sectional vs panel IC) and a data-vs-spec mismatch (5d horizon doesn't match where the signal lives). This is the standing review process working as intended — gates exist precisely so that 6-7 hour compute runs don't burn on shaky foundations. Granular session-by-session detail is at docs/sessions/larger_universe_v1/session_log.md. The diagnostic doc at docs/diagnostics/larger_universe_v1_horizon_diagnostic.md is the most current technical state. Mike will decide path forward tomorrow."),
    ]
    ref = ref_for_insert
    for style, text in items:
        ref = insert_paragraph_after(ref, text, style=style)

    # ----- Insertion 2: Add Appendix A pointer to session log.
    # Re-locate after the earlier inserts (paragraph indexes shifted).
    appendix_b_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("Appendix B"):
            appendix_b_idx = i
            break
    if appendix_b_idx is None:
        raise SystemExit("appendix B not found after earlier inserts")

    insert_after = d.paragraphs[appendix_b_idx - 1]
    appendix_items = [
        ("Heading 2", "Session log pointer — Larger Universe v1 study"),
        ("Normal",
         "Granular session-by-session activity log for the Larger Universe v1 study is at docs/sessions/larger_universe_v1/session_log.md. This is the canonical source of current detail; the tracker section above is a summary."),
    ]
    ref = insert_after
    for style, text in appendix_items:
        ref = insert_paragraph_after(ref, text, style=style)

    d.save(TRACKER)
    print(f"updated {TRACKER}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
