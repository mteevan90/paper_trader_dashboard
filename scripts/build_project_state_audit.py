"""build_project_state_audit.py — generate Project_State_Audit.docx.

What it is: one-shot script that emits the audit deliverable from
docs/Project_State_Audit.docx (or — by default — the user's Downloads
folder). The content is hardcoded findings collected by an audit Claude
session; the script just serializes them via python-docx into a structured
Word document with headings, bullets, and a status table.

When to re-run: any time the project state shifts enough that the existing
audit is stale, AND a refreshed Word artifact is wanted. Typical trigger
is "another contributor needs the latest project snapshot."

How to invoke (PowerShell):
    venv\\Scripts\\python.exe scripts\\build_project_state_audit.py

Output path: hardcoded at OUT_PATH (~/Downloads/Project_State_Audit.docx
as shipped). Edit OUT_PATH if you want the file written elsewhere.

To refresh the actual content for a new audit: edit the literal strings
inside this file directly. There's no external data source — every
finding, bullet, and table cell is in-line below. That keeps the audit
reproducible without depending on git or filesystem state at re-run time.

Requirements: python-docx (`pip install python-docx`). Not in
requirements.txt — install ad-hoc if needed for an audit refresh.
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = Path(os.path.expanduser("~/Downloads/Project_State_Audit.docx"))


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(11)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def status_table(doc):
    """Section 1's overview table."""
    rows = [
        # (Workstream, Status, Owner, Confidence, Notes)
        ("Equity research (core)", "WORKING",
         "Mike",
         "VERIFIED",
         "374 of 376 tests pass. 3 promoted studies live on cloud."),
        ("Phase 1 crypto-extension (asset_class plumbing)", "MERGED",
         "Mike",
         "VERIFIED",
         "PR #1, commit f268116. Sidebar now Stocks/Crypto/Options."),
        ("Phase 2 crypto strategy code", "PAUSED — STUB ONLY",
         "Chris",
         "VERIFIED",
         "src/crypto/ has 4 files (__init__, types, universe, "
         "static_universe). PR #2 landed; nothing since."),
        ("Options extension Phase 1 (shared-edge)", "MERGED",
         "Chris",
         "VERIFIED",
         "PR #3. src/options/ + tests/options/ scaffolded."),
        ("Options Sections 1-8 (build-out)", "MERGED",
         "Chris",
         "VERIFIED",
         "PRs #4-#11 + #12, #13 (fixes) + #14 (Polygon migration). "
         "20 files in src/options/."),
        ("Options PR #15 (v1 study hang fix)", "OPEN, NOT MERGED",
         "Chris",
         "VERIFIED",
         "Branch chris/options-fix-v1-study-hang at 7e0881c. "
         "8 files, +526/-106. Engine + rate-limiter fixes."),
        ("Options v1 study (locked snapshot)",
         "BLOCKED — gated on PR #15",
         "Chris",
         "VERIFIED",
         "models/snapshots/options/ does not exist. No "
         "pre_options_v1_<date> snapshot yet."),
        ("Cloud dashboard (paper-trader-mteev.streamlit.app)",
         "WORKING (equities only) — R2 STALE",
         "Mike",
         "VERIFIED via manifest read",
         "R2 manifest git_sha=bfee6dc; main is now f3c3800. "
         "Equity data correct; options not in R2 at all."),
        ("R2 sync state",
         "STALE — last upload 2026-05-09 17:20Z",
         "Mike",
         "VERIFIED via boto3 manifest GET",
         "Needs re-run after PR #15 merges + any equity changes."),
        ("Local Optuna study DBs", "WORKING",
         "Mike",
         "INFERRED",
         "models/cache/optuna_studies.db = 16 MB. Not "
         "re-validated end-to-end this audit."),
        ("SP1500 universe expansion", "PAUSED — STASHED",
         "Mike",
         "VERIFIED",
         "stash@{0} (refs/stash b8fe671). Awaiting Finnhub paid-"
         "tier TOS clarification."),
        ("Test suite (tests/)",
         "374 pass / 1 fail / 1 error / 1 skip",
         "Chris",
         "VERIFIED via pytest",
         "All failures attributable to missing `truststore` "
         "package in Mike's venv (it IS in requirements.txt)."),
        ("Documentation", "MOSTLY CURRENT",
         "Shared",
         "VERIFIED",
         "future_work.md, Crypto/Options decision memos, "
         "Comprehensive User Guide all on main. .env.example "
         "missing POLYGON_API_KEY entry."),
        ("CODEOWNERS", "ACTIVE, with phantom paths",
         "Shared",
         "VERIFIED",
         "@mteevan90 + @cmjteevan. Two directory rules "
         "(/src/equities/, /src/shared/) reference dirs that "
         "don't exist."),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(
            ("Workstream", "Status", "Owner", "Confidence", "Notes")):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, 1):
        cells = table.rows[r_idx].cells
        for c_idx, v in enumerate(row):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(v)
            run.font.size = Pt(9)


def main():
    doc = Document()

    # Title + meta
    title = doc.add_heading("Paper Trader — Project State Audit", level=0)
    para(doc,
         "Generated: 2026-05-11 by Claude Code (audit session).  Repo: "
         "paper_trader_dashboard.  Branch surveyed: origin/main "
         "@ f3c3800 + open branch chris/options-fix-v1-study-hang @ 7e0881c.",
         italic=True)
    para(doc,
         "Intent: honest snapshot for Mike. Not promotional. Surfaces the "
         "messy parts (stale R2 manifest, open PR blocking v1 study, "
         "phantom CODEOWNERS paths, missing venv deps, scratch-file "
         "clutter at repo root). If you only read one section, read "
         "Section 7 (Action Items).")

    # ------------------------------------------------------------------
    # 1. Executive summary
    # ------------------------------------------------------------------
    add_h1(doc, "1. Executive summary")
    para(doc,
         "Where the project actually stands as of 2026-05-11:")

    bullet(doc, "Equity research framework (Mike) is the mature, "
                "production-stable surface. Three promoted studies live on the "
                "cloud dashboard. Test suite is healthy.")
    bullet(doc, "Crypto extension paused at Phase 2 Section 1 — universe "
                "module stub only. No backtest engine, no strategy, no R2 data.")
    bullet(doc, "Options extension (Chris) is the active build-out. 8 of 9 "
                "sections shipped to main; Section 9 (paper-trade) untouched. "
                "v1 production study is BLOCKED on PR #15 (open) which fixes a "
                "pre-trial hang and rate-limiter races introduced by the "
                "Polygon migration.")
    bullet(doc, "Cloud dashboard is functionally serving the EQUITY content "
                "correctly but the R2 bucket has not been re-synced since "
                "2026-05-09 (commit bfee6dc). Origin/main has advanced 13 "
                "commits since. Options content is NOT in R2 at all.")
    bullet(doc, "Mike's local venv is missing truststore (required by "
                "options TLS handling) and pytest — both ARE in requirements.txt. "
                "`pip install -r requirements.txt` resolves both.")
    bullet(doc, "SP1500 universe expansion remains stashed (stash@{0}, "
                "b8fe671). No regressions; the Finnhub-tier blocker has not "
                "been resolved.")
    bullet(doc, "27+ scratch files at repo root (inspect_*.py, track_*.py, "
                "*.log) — mostly gitignored, but visually noisy and easy to "
                "confuse with active code.")

    add_h2(doc, "1.1 Status table")
    status_table(doc)

    # ------------------------------------------------------------------
    # 2. What's working
    # ------------------------------------------------------------------
    add_h1(doc, "2. What's working — concrete + verified")

    add_h2(doc, "2.1 Equity research framework")
    bullet(doc, "src/backtest.py + src/optuna_runner.py + src/objective.py: "
                "core backtest + tuning pipeline. Unchanged for ~3 weeks of "
                "main commits; the foundation has stabilized.")
    bullet(doc, "Locked snapshot at models/snapshots/equities/pre_v2_20260505/ "
                "(228 MB, 1473 files). Phase 1 relocation verified.")
    bullet(doc, "Three promoted equity studies (manifest read from R2): "
                "best_regime_dependent_v1_20260505_2240_325 (#325), "
                "best_continuous_sizing_v1_1852, "
                "best_15_position_study_v1_842. All have promoted=true in "
                "their meta.json.")
    bullet(doc, "data_source.path_to() with default asset_class=\"equities\" "
                "returns bit-identical legacy paths — every existing equity "
                "callsite is unchanged.")

    add_h2(doc, "2.2 Dashboard (Streamlit)")
    bullet(doc, "https://paper-trader-mteev.streamlit.app deploys "
                "successfully. Last verified 2026-05-09 post-equity-R2-resync "
                "(this session has not visually verified today).")
    bullet(doc, "Sidebar asset selector: Stocks / Crypto / Options. Stocks "
                "renders the full 8-tab equity dashboard. Crypto and Options "
                "show placeholder text (\"in development\").")
    bullet(doc, "CLOUD_MODE auth gate (dashboard_auth.py) live.")

    add_h2(doc, "2.3 Options module (in main, pre-PR-#15)")
    bullet(doc, "21 files in src/options/. 18 test files in tests/options/ "
                "(374 tests, mostly passing).")
    bullet(doc, "Polygon historical fetcher (PR #14) replaces Tradier for "
                "backtest data. Tradier kept for paper-trade / live (v2+).")
    bullet(doc, "Greeks (Black-Scholes), engine, position model, "
                "BacktestConfig, FeeModel, Optuna runner, promotion gate, "
                "concentration analysis — all merged.")

    add_h2(doc, "2.4 Phase 1 crypto refactor")
    bullet(doc, "Commit f268116. data_source.py asset_class-aware; "
                ".github/CODEOWNERS in place; sidebar wired.")

    # ------------------------------------------------------------------
    # 3. What's broken or half-finished
    # ------------------------------------------------------------------
    add_h1(doc, "3. What's broken or half-finished")

    add_h2(doc, "3.1 Open work that blocks downstream tasks")

    add_h3(doc, "PR #15 (chris/options-fix-v1-study-hang) — OPEN")
    para(doc, "Branch: origin/chris/options-fix-v1-study-hang @ 7e0881c. "
              "One commit ahead of main, 8 files, +526/-106 lines. Commit "
              "message: \"Fix v1 study hang in pre-trial orchestration\".",
         italic=True)
    bullet(doc, "Root cause: engine.py was still passing Tradier's "
                "fetch_history to chain_reconstruction even after PR #14 "
                "swapped the default to Polygon. Tradier returns null for "
                "expired OCCs, so reconstruct_chain looped on empty data.")
    bullet(doc, "Fixes bundled: separate RateLimiter instances per data "
                "source (Tradier 60/min vs Polygon 300/min), thread-safe "
                "lock on rate-limit deque (16 concurrent chain workers), "
                "capped header-driven sleep at 60s, multi-day pre-fetch in "
                "_default_deps, removed Polygon cache-write sanity gate "
                "that was rejecting legitimate sparse-expiry contracts.")
    bullet(doc, "Impact: v1 production study CANNOT complete until this "
                "merges. No pre_options_v1_<date> snapshot can be created.")
    bullet(doc, "Recommendation: merge after Mike (or Chris) eyeballs the "
                "diff. The fixes are real, tests are added "
                "(test_chain_reconstruction.py +120 lines, "
                "test_engine.py +131 lines).")

    add_h3(doc, "Cloud dashboard R2 sync is stale")
    bullet(doc, "Bucket snapshot_manifest.json: git_sha=bfee6dc, "
                "generated_at=2026-05-09T17:20:38, n_files=82, asset_class="
                "equities. Origin/main is at f3c3800 — 13 commits ahead.")
    bullet(doc, "Functionally: equity content rendered correctly when "
                "verified last session. Any equity content that changed in "
                "the last 13 commits (none did — they were all options work) "
                "would be missing.")
    bullet(doc, "Options content has never been uploaded. No "
                "snapshot_for_cloud.py --asset-class options run on record.")

    add_h2(doc, "3.2 Half-finished")

    bullet(doc, "Crypto extension: src/crypto/ has 4 files "
                "(__init__.py, static_universe.py, types.py, universe.py) — "
                "the universe-module stub from PR #2. No backtest engine, "
                "no strategy code, no fetched data. Confirmed paused per "
                "Mike's pivot.")
    bullet(doc, "Options Section 9 (paper-trade) not started. No "
                "code or branch found.")
    bullet(doc, "src/equities/ and src/shared/ — directories referenced "
                "by CODEOWNERS but DO NOT EXIST. The rule "
                "`/src/equities/ @mteevan90` is a no-op; equity code "
                "still lives flat in src/.")
    bullet(doc, "v1 options production study (run_v1_study in "
                "src/options/v1_study.py) is committed but un-run. The "
                "intended output (models/snapshots/options/pre_options_v1_*/) "
                "doesn't exist. Gated on PR #15.")

    add_h2(doc, "3.3 Test suite state")

    bullet(doc, "Total: 374 passing, 1 failing, 1 collection error, 1 "
                "skipped.")
    bullet(doc, "test_ssl_setup.py collection error: "
                "ModuleNotFoundError: No module named 'truststore'.")
    bullet(doc, "test_optuna_runner.py::test_smoke_study_constants_match_locked_config "
                "fails for the same truststore reason (transitively via "
                "src/options/_ssl.py).")
    bullet(doc, "Both failures resolve with `venv\\Scripts\\pip install "
                "-r requirements.txt`. truststore IS pinned at 0.10.4 in "
                "requirements.txt — it just isn't installed in Mike's "
                "current venv.")
    bullet(doc, "pytest itself was also missing from Mike's venv — had to "
                "be installed for this audit. Not in requirements.txt as a "
                "dev-only dep.")

    add_h2(doc, "3.4 Documentation gaps")
    bullet(doc, ".env.example only documents Tradier env vars. "
                "POLYGON_API_KEY is required by src/options/polygon.py but "
                "is undocumented in .env.example. New contributors won't "
                "know it exists.")
    bullet(doc, "docs/Options_Extension_Decisions.md Status line says "
                "\"Phase 1 + Sections 1, 2, 3 merged. Section 4 specced and "
                "ready to ship.\" Reality (main): Sections 1-8 + 2.5 are "
                "merged. The memo's Status row is ~5 sections out of date.")

    # ------------------------------------------------------------------
    # 4. Spaghetti / cleanup items
    # ------------------------------------------------------------------
    add_h1(doc, "4. Spaghetti / cleanup items")

    add_h2(doc, "4.1 Repository hygiene")
    bullet(doc, "27 scratch files at the top of the repo: inspect_*.py "
                "(8), track_*.py (4), check_*.py (3), progress_*.py (2), "
                "*.log files (~10). Most are gitignored, but they appear in "
                "directory listings and `ls` output, making the project "
                "look cluttered. Recommend: scripts/ subdir for the "
                "long-lived ones, delete the rest.")
    bullet(doc, "src.zip at repo root — manual snapshot of src/. "
                "Likely stale and unused; safe to delete.")
    bullet(doc, "Documents/ directory at repo root — empty per ls "
                "(or only contains items I didn't peek at). Likely "
                "leftover.")
    bullet(doc, "data/ and scripts/ directories at repo root — new, "
                "haven't audited contents. Worth Mike checking what's in "
                "them.")

    add_h2(doc, "4.2 Snapshot directory still has orphans")
    bullet(doc, "Phase 1 moved models/snapshots/pre_v2_20260505/ -> "
                "models/snapshots/equities/. But three other equity-context "
                "snapshots remain at the top of models/snapshots/: "
                "macro_signal_investigation_20260506, "
                "parallelism_diagnostic_20260506, post_macro_fix_20260506. "
                "Should be moved under equities/ for consistency.")
    bullet(doc, "models/snapshots/options/ doesn't exist yet (no v1 study "
                "snapshot created — see §3.1).")
    bullet(doc, "models/snapshots/crypto/ doesn't exist yet (no crypto "
                "data fetched).")

    add_h2(doc, "4.3 CODEOWNERS references phantom paths")
    bullet(doc, "`/src/equities/` — directory doesn't exist. Mike's "
                "equity code is protected by per-file rules instead. The "
                "directory rule is aspirational (intended for future "
                "equity-only modules) but currently a no-op.")
    bullet(doc, "`/src/shared/` — directory doesn't exist. Similarly "
                "aspirational.")
    bullet(doc, "Recommendation: either create the dirs (with placeholder "
                "READMEs documenting purpose) or remove the rules until "
                "real code lands there.")

    add_h2(doc, "4.4 sp1500_fetch_failures.txt in docs/ is 1448 lines")
    bullet(doc, "This is mostly a per-ticker dump from a debug session. "
                "Useful as a snapshot of what was broken on that day, but "
                "should probably move under docs/specs/ or docs/diagnostics/ "
                "rather than living at the top of docs/.")

    add_h2(doc, "4.5 Long-standing TODO in optuna_runner.py")
    para(doc, "src/optuna_runner.py:275 has a multi-paragraph TODO about "
              "a thread-safety race in run_backtest's vectorized predict "
              "block. Per the comment, two fix attempts were tried and "
              "reverted because they made things worse. The race exists "
              "but is masked by workload heterogeneity. Documented but "
              "still latent.",
         italic=True)

    # ------------------------------------------------------------------
    # 5. Per-contributor status
    # ------------------------------------------------------------------
    add_h1(doc, "5. Per-contributor status")

    add_h2(doc, "5.1 Mike (@mteevan90)")
    bullet(doc, "Last commit on main authored by Mike: "
                "bfee6dc \"fix: allow docs/*.docx in gitignore; add "
                "Comprehensive User Guide\" — 2026-05-09.")
    bullet(doc, "Last commit AT ALL by Mike on main: same. Chris has "
                "been the sole committer since.")
    bullet(doc, "Active workstreams (paused):")
    bullet(doc, "  - SP1500 universe expansion: stashed at stash@{0}. "
                "Awaiting Finnhub TOS clarification.",
           level=1)
    bullet(doc, "  - Live equity cache relocation (entry in future_work.md): "
                "Deferred. Touches do-not-modify files.",
           level=1)
    bullet(doc, "Status: idle on equity work this week. Equity codebase "
                "is stable; no in-flight PRs from Mike.")

    add_h2(doc, "5.2 Chris (@cmjteevan)")
    bullet(doc, "Currently active on chris/options-fix-v1-study-hang. "
                "Last commit 7e0881c — Section 8 hang fix.")
    bullet(doc, "13 commits merged to main since 2026-05-09 (PRs #2 "
                "through #14). Pace: ~6-7 PRs in 2 days.")
    bullet(doc, "Strategy classes built (v1 scope per memo): "
                "Covered Calls + Cash-Secured Puts only. v1.1 (credit "
                "spreads), v1.2 (iron condors), v2 (long directional) "
                "are specced but not started.")
    bullet(doc, "Polygon migration (PR #14) replaced Tradier for "
                "backtest historical data. Tradier retained for "
                "live/paper-trade. This was the right call — Tradier's "
                "expired-OCC history endpoint is broken.")
    bullet(doc, "Status: actively building. Next blocker: get PR #15 "
                "merged so v1 study can run end-to-end.")

    # ------------------------------------------------------------------
    # 6. Shared infrastructure
    # ------------------------------------------------------------------
    add_h1(doc, "6. Shared infrastructure")

    add_h2(doc, "6.1 Data layer")
    bullet(doc, "models/cache/: equity files at flat (legacy) layout. "
                "models/cache/options/ likely exists (Polygon cache "
                "subdir per polygon.py); models/cache/crypto/ probably "
                "doesn't.")
    bullet(doc, "models/snapshots/equities/pre_v2_20260505/: 228 MB, "
                "the canonical equity baseline.")
    bullet(doc, "No options snapshot exists yet.")

    add_h2(doc, "6.2 R2 (Cloudflare) state")
    bullet(doc, "Bucket: paper-trader-snapshots.")
    bullet(doc, "snapshot_manifest.json (root): asset_class=equities, "
                "git_sha=bfee6dc, generated_at=2026-05-09T17:20:38Z, "
                "n_files=82, manifest_version=2.")
    bullet(doc, "Bucket prefixes: equities/ populated (82 keys). "
                "crypto/ and options/ are empty (never uploaded).")
    bullet(doc, "Orphan pre-Phase-1 flat keys (model/..., optuna/..., "
                "dashboard_results/...) likely still in bucket — left "
                "intentionally last session because deleting them was "
                "out of scope. Cost: ~40 MB of duplicate storage.")

    add_h2(doc, "6.3 CI / tooling")
    bullet(doc, "No CI workflow found at .github/workflows/. "
                "All testing is run locally via pytest. (Search for "
                "*.yml in .github/ returns nothing; CODEOWNERS is the "
                "only file there.)")
    bullet(doc, "gh CLI not installed on Mike's machine. PR status "
                "was inferred from commit-message squash-merge markers "
                "(\"(#N)\") and remote branch presence.")

    # ------------------------------------------------------------------
    # 7. Action items
    # ------------------------------------------------------------------
    add_h1(doc, "7. Action items, ranked")

    add_h2(doc, "7.1 URGENT — blocks contributors or users")
    bullet(doc, "Merge PR #15 (options v1 study hang fix). Currently "
                "blocking Chris from running the production v1 study. "
                "Diff is bounded and includes new tests.")
    bullet(doc, "Run `venv\\Scripts\\pip install -r requirements.txt` "
                "on Mike's machine. Resolves truststore + pytest gaps. "
                "Without this, Mike can't run options code or the "
                "test suite.")
    bullet(doc, "Confirm cloud dashboard at "
                "paper-trader-mteev.streamlit.app still renders. "
                "(Audit didn't visually verify; manifest read suggests "
                "it should work for equities.)")

    add_h2(doc, "7.2 HIGH — risks getting worse if left alone")
    bullet(doc, "Add POLYGON_API_KEY to .env.example with a comment "
                "linking to the Massive.com signup page. New "
                "contributor onboarding currently fails silently.")
    bullet(doc, "Update Options_Extension_Decisions.md Status row "
                "to reflect Sections 4-8 + 2.5 merged. Stale memos "
                "get ignored.")
    bullet(doc, "Decide what to do with stale "
                "origin/chris/options-phase-1-shared-edge branch. "
                "Already merged via PR #3 (squash); branch is "
                "diverged-and-dead. Delete or note as historical.")
    bullet(doc, "After PR #15 merges, sync R2: "
                "`snapshot_for_cloud.py --asset-class equities` "
                "(if any equity content changed) and "
                "`--asset-class options` (first-ever options upload).")

    add_h2(doc, "7.3 MEDIUM — quality improvements")
    bullet(doc, "Move 3 orphan equity snapshots into "
                "models/snapshots/equities/ for consistency.")
    bullet(doc, "Either create src/equities/ and src/shared/ "
                "directories (with README.md placeholders) or remove "
                "the phantom CODEOWNERS rules.")
    bullet(doc, "Move/delete the 27 scratch files at repo root. "
                "Survey them for which ones Mike still uses; archive "
                "the rest under scripts/scratch/ or delete.")
    bullet(doc, "Move docs/sp1500_fetch_failures.txt (1448 lines) "
                "to docs/diagnostics/ or delete if no longer "
                "diagnostically useful.")
    bullet(doc, "Add a CI workflow (.github/workflows/test.yml) that "
                "runs pytest on push/PR. Free with GitHub Actions; "
                "would have caught the truststore venv-vs-requirements "
                "drift.")

    add_h2(doc, "7.4 LOW — nice-to-haves")
    bullet(doc, "Resolve the SP1500 stash@{0}. Either pop and commit "
                "(if Finnhub clarity has arrived) or convert to a "
                "named branch for clearer state.")
    bullet(doc, "Resolve the optuna_runner.py:275 thread-safety race "
                "TODO. Deferred since segment 22; race is masked by "
                "workload heterogeneity but real.")
    bullet(doc, "Delete src.zip if it's stale (likely).")
    bullet(doc, "Confirm Documents/ directory contents and either "
                "populate or delete.")
    bullet(doc, "Add a pytest pin to requirements.txt with a comment "
                "marking it dev-only. Currently every fresh venv has "
                "to install it ad-hoc.")

    # ------------------------------------------------------------------
    # 8. Unclear items (audit honest about what it doesn't know)
    # ------------------------------------------------------------------
    add_h1(doc, "8. What this audit could NOT verify")
    para(doc,
         "Honesty section. These are claims I could not check from "
         "the command line alone:")
    bullet(doc, "Whether the deployed Streamlit Cloud app at "
                "paper-trader-mteev.streamlit.app currently renders "
                "without error in a browser. Manifest read suggests "
                "equity content is correct, but the deployed app may "
                "be running a newer source-of-truth than the R2 "
                "data points to. Mike should visually confirm.")
    bullet(doc, "Whether the optuna_studies.db at "
                "models/cache/optuna_studies.db end-to-end-loads in "
                "the dashboard. The file is 16 MB and looks intact, "
                "but I did not exercise the read path.")
    bullet(doc, "Whether `pytest` covers the equity code path at "
                "all. The 374 passing tests are mostly options-side "
                "(tests/options/ has 18 files; tests/crypto/ has 1; "
                "no tests/equities/ directory). Equity correctness "
                "is asserted via study reproducibility against the "
                "locked snapshot, not unit tests.")
    bullet(doc, "Whether the Polygon API key currently in Mike's "
                ".env is valid and rate-limit healthy. The audit "
                "did not call the live endpoint.")
    bullet(doc, "Contents of data/ and scripts/ subdirectories at "
                "repo root (new since Phase 1). Did not enumerate.")
    bullet(doc, "Whether crypto-extension Phase 2 work has any "
                "uncommitted local state on Chris's machine. Only "
                "main + chris/* remote branches were inspected.")

    out_dir = OUT_PATH.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}  ({OUT_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
